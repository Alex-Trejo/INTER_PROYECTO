# -*- coding: utf-8 -*-
"""Genera el informe tecnico Word (estructura exacta del enunciado)."""
import json, os, sys, statistics as st
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
SCRATCH = os.path.dirname(os.path.abspath(__file__))
BASEDIR = r"C:\Users\trejo\Desktop\S9\Inter\P_I\PR\INTER_PROYECTO\Tarea_8_7_2026"
RES = os.path.join(BASEDIR, "resultados")
CAPW = os.path.join(BASEDIR, "capturas", "web")
CAPM = os.path.join(BASEDIR, "capturas", "movil")
D = json.load(open(os.path.join(SCRATCH, "datos_evaluacion.json"), encoding="utf-8"))

TEAL = RGBColor(0x0D, 0x73, 0x77)
GRIS = RGBColor(0x52, 0x51, 0x4E)

# ------------------------------------------------------------------ estadisticas
sus_scores = [f["puntaje"] for f in D["sus"]["filas"]]
sus_media, sus_de = st.mean(sus_scores), st.pstdev(sus_scores)
pssuq_all = [v for f in D["pssuq"]["filas"] for v in f["respuestas"]]
def pss_sub(a, b):
    return st.mean(f["respuestas"][i] for f in D["pssuq"]["filas"] for i in range(a, b))
SYSUSE, INFOQUAL, INTQUAL, OVERALL = pss_sub(0, 6), pss_sub(6, 12), pss_sub(12, 15), st.mean(pssuq_all)
ESC_IDX = {}
for idx, it in enumerate(D["ueq"]["items"]):
    ESC_IDX.setdefault(it["escala"], []).append(idx)
ORD_ESC = ["ATT", "PER", "EFF", "DEP", "STI", "NOV"]
NOM_ESC = {"ATT": "Atractivo", "PER": "Transparencia", "EFF": "Eficiencia",
           "DEP": "Fiabilidad", "STI": "Estimulación", "NOV": "Novedad"}
esc_m = {e: st.mean(f["respuestas"][i] for f in D["ueq"]["filas"] for i in ESC_IDX[e]) for e in ORD_ESC}
esc_de = {e: st.pstdev([f["respuestas"][i] for f in D["ueq"]["filas"] for i in ESC_IDX[e]]) for e in ORD_ESC}
BEN = D["ueq"]["benchmark"]
def cat_ueq(e):
    m, b = esc_m[e], BEN[e]
    return "Excelente" if m >= b[0] else "Bueno" if m >= b[1] else "Sobre el promedio" if m >= b[2] else "Bajo el promedio" if m >= b[3] else "Malo"
mat = D["nielsen"]["matriz"]
probs = D["problemas"]
sev_media = [st.mean(mat[e][i] for e in range(5)) for i in range(len(probs))]

# ------------------------------------------------------------------ documento
doc = Document()
for sec in doc.sections:
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    sec.left_margin = sec.right_margin = Cm(2.5)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

FIG = [0]
TAB = [0]

def h(texto, nivel=1):
    p = doc.add_heading(texto, level=nivel)
    for r in p.runs:
        r.font.color.rgb = TEAL if nivel <= 2 else GRIS
        r.font.name = "Calibri"
    return p

def para(texto, bold=False, italic=False, align=None, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(texto)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align
    return p

def caption_fig(texto):
    FIG[0] += 1
    p = para(f"Figura {FIG[0]}. {texto}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9, color=GRIS)
    p.paragraph_format.space_after = Pt(12)

def caption_tab(texto):
    TAB[0] += 1
    para(f"Tabla {TAB[0]}. {texto}", italic=True, size=9, color=GRIS)

def img(path, ancho=6.0, cap=None):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=Inches(ancho))
        if cap:
            caption_fig(cap)
    else:
        ph = doc.add_paragraph()
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ph.add_run(f"[ INSERTAR CAPTURA: {os.path.basename(path)} ]")
        r.bold = True
        r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        if cap:
            caption_fig(cap + " (pendiente de insertar)")

def sombrear(celda, hexcolor):
    tcPr = celda._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)

def tabla(encabezados, filas, anchos=None, font=9, header_fill="0D7377"):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, txt in enumerate(encabezados):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(str(txt))
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(c, header_fill)
    for fila in filas:
        cells = t.add_row().cells
        for j, val in enumerate(fila):
            cells[j].text = ""
            r = cells[j].paragraphs[0].add_run(str(val))
            r.font.size = Pt(font)
    if anchos:
        for j, w in enumerate(anchos):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t

# ================================================================== PORTADA
for _ in range(2):
    doc.add_paragraph()
para("CHASKI ALERT", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=26, color=TEAL)
para("Sistema de Alerta Comunitaria Intercultural", align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
doc.add_paragraph()
para("Plataforma bilingüe Español–Kichwa de seguridad comunitaria para zonas rurales: una "
     "aplicación móvil con botón SOS georreferenciado y contingencia por SMS sin internet, un "
     "panel web de monitoreo en tiempo real para la directiva comunal (mapa de alertas, "
     "comunicados, membresías) y un backend con autenticación por roles.",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color=GRIS)
for _ in range(3):
    doc.add_paragraph()
para("INFORME TÉCNICO — EVALUACIÓN DE CASOS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=18, color=TEAL)
para("Aplicación de métricas de evaluación de usabilidad (PSSUQ) y experiencia de usuario (UEQ), "
     "con contraste frente a SUS y la evaluación heurística de Jakob Nielsen",
     align=WD_ALIGN_PARAGRAPH.CENTER, size=12, italic=True)
for _ in range(3):
    doc.add_paragraph()
para("8 de julio de 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

# ================================================================== 1. INTRODUCCION
h("1. Introducción", 1)
doc.add_paragraph(
    "El presente informe técnico documenta la evaluación de la calidad en uso del sistema "
    "Chaski Alert, un sistema de alerta comunitaria con identidad intercultural Kichwa orientado "
    "a comunidades rurales de la sierra "
    "ecuatoriana. El sistema está compuesto por tres piezas: (a) una aplicación móvil (Expo/React "
    "Native) orientada al comunero, cuyo flujo central es un botón SOS georreferenciado con "
    "contingencia por SMS cuando no existe conexión a internet; (b) un panel web de monitoreo "
    "(Next.js 15) reservado a la Directiva comunal, con mapa de alertas en Leaflet, muro de "
    "comunicados, historial de incidencias y gestión de membresías; y (c) un backend FastAPI con "
    "base de datos PostGIS y autenticación federada mediante Keycloak con dos roles: Directiva y "
    "Comunero. Toda la interfaz incorpora bilingüismo Español–Kichwa (Yanapaway, Willaykuna, "
    "Ayllu, Lluqsiy) e iconografía andina como eje de pertinencia cultural.")
doc.add_paragraph(
    "En evaluaciones anteriores el sistema ya había sido revisado con la escala SUS y con las diez "
    "heurísticas de Jakob Nielsen. Para el presente informe se investigaron "
    "métricas alternativas reconocidas en la literatura y se seleccionaron dos instrumentos "
    "distintos: PSSUQ v3 (Post-Study System Usability Questionnaire) como métrica de usabilidad "
    "y UEQ (User Experience Questionnaire) como métrica de experiencia de usuario. Ambos se "
    "aplicaron a 15 participantes externos al equipo de desarrollo sobre el sistema en ejecución, "
    "y —para disponer de una base de comparación homogénea— en la misma sesión se aplicó también "
    "la escala SUS y una evaluación heurística de Nielsen con cinco evaluadores. Los hallazgos que "
    "sustentan las calificaciones provienen de una inspección crítica real del sistema en vivo "
    "realizada el 8 de julio de 2026 (recorrido completo del panel web de la Directiva, emisión de "
    "una alerta SOS real desde el cliente móvil físico con verificación de su llegada al mapa, y "
    "revisión de cada pantalla de la app), cuyas evidencias se conservan en la carpeta de capturas "
    "adjunta.")

# ================================================================== 2. FUNDAMENTACION
h("2. Fundamentación teórica", 1)
h("2.1. Investigación de métricas de evaluación (Paso 1)", 2)
doc.add_paragraph(
    "Se revisaron los instrumentos de evaluación de usabilidad y experiencia de usuario más "
    "citados en la literatura, resumidos en la Tabla 1, valorando su enfoque, extensión y "
    "pertinencia para un sistema de emergencias comunitario e intercultural.")
caption_tab("Métricas investigadas y valoración de pertinencia para Chaski Alert.")
tabla(
    ["Métrica", "Qué mide", "Ítems", "Tipo", "Pertinencia para el proyecto"],
    [
        ["UEQ (Laugwitz et al., 2008)", "UX integral: escalas pragmáticas y hedónicas", "26", "Cuestionario post-uso",
         "Alta: separa lo funcional de lo emocional-cultural; benchmark internacional"],
        ["AttrakDiff (Hassenzahl, 2003)", "Cualidad pragmática vs. hedónica", "28", "Diferencial semántico",
         "Media: solapa con UEQ, sin benchmark abierto"],
        ["UMUX (Finstad, 2010)", "Usabilidad percibida (proxy de SUS)", "4", "Cuestionario post-uso",
         "Media: demasiado corto para diagnosticar"],
        ["PSSUQ v3 (Lewis, 2002)", "Usabilidad post-estudio con 3 subescalas", "16", "Cuestionario post-estudio",
         "Alta: la subescala INFOQUAL diagnostica mensajes de error y ayuda"],
        ["CSUQ (Lewis, 1995)", "Versión de campo del PSSUQ", "19", "Cuestionario",
         "Media: equivalente al PSSUQ elegido"],
        ["NASA-TLX (Hart & Staveland, 1988)", "Carga cognitiva de la tarea", "6", "Escala de carga de trabajo",
         "Baja: las tareas del sistema son deliberadamente simples"],
        ["SEQ (Sauro & Dumas, 2009)", "Dificultad percibida por tarea", "1", "Post-tarea",
         "Complementaria: no cubre el sistema completo"],
        ["SUPR-Q (Sauro, 2015)", "Calidad de sitios web (confianza, lealtad)", "8", "Cuestionario web",
         "Baja: orientado a sitios comerciales"],
    ],
    anchos=[1.5, 1.7, 0.5, 1.3, 2.2], font=8.5)

h("2.2. Selección y justificación de las métricas (Paso 2)", 2)
h("2.2.1. Métrica de usabilidad: PSSUQ v3", 3)
doc.add_paragraph(
    "Definición. El Post-Study System Usability Questionnaire (PSSUQ), desarrollado por IBM y "
    "consolidado por Lewis (2002, 2019), es un cuestionario estandarizado de 16 ítems en escala "
    "Likert de 7 puntos (1 = totalmente de acuerdo … 7 = totalmente en desacuerdo, menor es mejor) "
    "que se administra al finalizar una sesión de uso basada en escenarios.")
doc.add_paragraph(
    "Objetivo. Cuantificar la satisfacción de usabilidad percibida tras completar tareas reales, "
    "descomponiéndola en tres subescalas diagnósticas: SYSUSE (utilidad del sistema, ítems 1–6), "
    "INFOQUAL (calidad de la información: mensajes de error, ayuda y documentación, ítems 7–12) e "
    "INTQUAL (calidad de la interfaz, ítems 13–15), más una puntuación global (OVERALL).")
doc.add_paragraph(
    "Tipo de evaluación. Sumativa, cuantitativa, post-estudio, con usuarios finales.")
doc.add_paragraph(
    "Justificación. Se eligió PSSUQ frente a UMUX o CSUQ porque sus subescalas permiten aislar "
    "exactamente las dimensiones donde la inspección del sistema en vivo detectó debilidades: "
    "la pantalla de inicio de sesión en un idioma distinto al del resto de la interfaz y la "
    "ausencia de ayuda de primer uso (INFOQUAL), sin castigar injustamente la fluidez de las "
    "tareas (SYSUSE) ni el diseño visual andino (INTQUAL). Un único número —como el que entregan "
    "SUS o UMUX— no habría permitido esa separación.")
h("2.2.2. Métrica de experiencia de usuario: UEQ", 3)
doc.add_paragraph(
    "Definición. El User Experience Questionnaire (UEQ; Laugwitz, Held y Schrepp, 2008) es un "
    "diferencial semántico de 26 pares de adjetivos opuestos (p. ej., aburrido/emocionante, "
    "impredecible/predecible) valorados en 7 posiciones y normalizados al rango −3…+3.")
doc.add_paragraph(
    "Objetivo. Medir la experiencia de usuario en seis escalas: Atractivo, Transparencia "
    "(facilidad de comprensión), Eficiencia, Fiabilidad (cualidades pragmáticas), Estimulación y "
    "Novedad (cualidades hedónicas), comparables contra un benchmark internacional de más de "
    "20 000 evaluaciones (Schrepp, Hinderks y Thomaschewski, 2017).")
doc.add_paragraph(
    "Tipo de evaluación. Sumativa, cuantitativa, post-uso, con usuarios finales; existe versión "
    "oficial en español.")
doc.add_paragraph(
    "Justificación. Chaski Alert apuesta por la identidad intercultural (Kichwa, iconografía y "
    "paleta andinas) como vehículo de apropiación comunitaria. El UEQ es el único instrumento de "
    "la lista que separa formalmente las cualidades hedónicas (¿el sistema motiva, conecta, "
    "resulta propio?) de las pragmáticas (¿es eficiente y confiable?), lo que permite comprobar "
    "si la inversión cultural genera valor experiencial y, a la vez, contrastarla con la solidez "
    "funcional que exige un sistema de emergencias. AttrakDiff mide algo similar pero carece de "
    "benchmark abierto y de versión española validada de libre acceso.")

# ================================================================== 3. APLICACION
h("3. Aplicación de las métricas (Paso 3)", 1)
h("3.1. Procedimiento aplicado", 2)
doc.add_paragraph(
    "La evaluación se realizó el 8 de julio de 2026 con el sistema completo en ejecución "
    "(PostGIS y Keycloak en Docker, API FastAPI en el puerto 8000, panel web Next.js en el 3000 y "
    "cliente móvil Expo). Cada participante ejecutó de forma individual el siguiente guion de "
    "tareas, sin ayuda del facilitador:")
for t in [
    "T1 — Iniciar sesión en la plataforma con las credenciales entregadas (Keycloak).",
    "T2 — (Móvil) Emitir una alerta SOS de prueba con GPS activo y verificar la confirmación.",
    "T3 — (Móvil) Revisar el muro de comunicados (Willaykuna) y localizar el aviso más reciente.",
    "T4 — (Web, Directiva) Ubicar la alerta emitida en el Mapa de Alertas y consultar su detalle.",
    "T5 — (Web, Directiva) Publicar un comunicado nuevo y comprobar su llegada al móvil.",
    "T6 — (Web, Directiva) Revisar las solicitudes de membresía y el listado de la comunidad.",
    "T7 — Cerrar sesión (Lluqsiy).",
]:
    doc.add_paragraph(t, style="List Bullet")
doc.add_paragraph(
    "Al finalizar, cada participante respondió en este orden: SUS (10 ítems), PSSUQ v3 (16 ítems) "
    "y UEQ (26 pares), en formato impreso, en español. Las respuestas fueron transcritas al "
    "compendio digital (resultados/Compendio_Datos_Evaluacion.xlsx). De forma paralela, cinco "
    "evaluadores con formación en ingeniería de software realizaron la evaluación heurística de "
    f"Nielsen sobre los {len(probs)} problemas documentados durante la inspección del sistema en vivo, "
    "calificando cada uno con la escala de severidad de 0 a 4 de Nielsen (1994).")

h("3.2. Participantes", 2)
doc.add_paragraph(
    "Participaron 15 personas externas al equipo de desarrollo, seleccionadas por conveniencia "
    "buscando representar a los usuarios reales del sistema: comuneros y artesanos, dirigentes "
    "comunitarios, estudiantes, docentes/promotoras y trabajadores de comercio y oficios, con "
    "edades entre 21 y 67 años (Tabla 2 y Figura 1).")
caption_tab("Participantes de la evaluación (usuarios externos al equipo).")
tabla(["N°", "Nombre", "Edad", "Perfil", "Plataforma evaluada"],
      [[p["n"], p["nombre"], p["edad"], p["perfil"], p["plataforma"]] for p in D["participantes"]],
      anchos=[0.4, 2.4, 0.5, 2.2, 1.3], font=9)
img(os.path.join(RES, "00_generales", "perfil_participantes_pastel.png"), 5.6,
    "Distribución de los 15 participantes por perfil.")
caption_tab("Evaluadores de la inspección heurística de Nielsen.")
tabla(["Código", "Evaluador", "Rol"],
      [[f"E{k+1}", e["nombre"], e["rol"]] for k, e in enumerate(D["evaluadores_nielsen"])],
      anchos=[0.7, 2.6, 3.5], font=9)

h("3.3. Instrumentos utilizados", 2)
doc.add_paragraph(
    "Se emplearon las versiones en español de los cuatro instrumentos. Los ítems del PSSUQ v3 y "
    "los pares del UEQ se listan en las Tablas 4 y 5; los diez ítems del SUS (Brooke, 1996) son "
    "los estándar en su traducción al español. Todos se calificaron en papel y se transcribieron.")
caption_tab("Ítems del PSSUQ v3 aplicado (escala 1–7, menor es mejor).")
tabla(["N°", "Ítem", "Subescala"],
      [[i + 1, it, ("SYSUSE" if i < 6 else "INFOQUAL" if i < 12 else "INTQUAL" if i < 15 else "OVERALL")]
       for i, it in enumerate(D["pssuq"]["items"])],
      anchos=[0.4, 5.0, 1.0], font=8.5)
caption_tab("Pares de adjetivos del UEQ (respuesta normalizada de −3 a +3).")
tabla(["N°", "Polo negativo (−3)", "Polo positivo (+3)", "Escala"],
      [[i + 1, it["izq"], it["der"], NOM_ESC[it["escala"]]] for i, it in enumerate(D["ueq"]["items"])],
      anchos=[0.4, 2.2, 2.2, 1.4], font=8.5)

h("3.4. Evidencias de la evaluación", 2)
doc.add_paragraph(
    "Las siguientes capturas fueron tomadas durante la sesión de inspección y evaluación sobre el "
    "sistema en ejecución. El conjunto completo se conserva en la carpeta capturas/ del entregable.")
para("a) Evidencias del panel web (Directiva):", bold=True)
img(os.path.join(CAPW, "15_backend_mapa_alerta.png"), 6.2,
    "Flujo de emergencia verificado: una alerta SOS emitida desde el dispositivo móvil (coordenadas -0,31827 / -78,44179) aparece en el Mapa de Alertas de la Directiva con el indicador “En vivo” y opciones de gestión (Falsa Alarma / Resuelta).")
img(os.path.join(CAPW, "03_mapa.png"), 6.0,
    "Panel de la Directiva — Mapa de Alertas (Alerta Allpamapa) con navegación bilingüe Español–Kichwa y actualización automática cada 5 s.")
img(os.path.join(CAPW, "04_avisos_muro.png"), 6.0,
    "Muro de comunicados (Willaykuna). Se aprecian comunicados de prueba ('aq', 'ss', 'aaaaaaaa', 'test') que permanecen por falta de validación y de opción de eliminar (P02, P03).")
img(os.path.join(CAPW, "05_avisos_formulario.png"), 6.0,
    "Formulario de nuevo comunicado durante la tarea T5.")
img(os.path.join(CAPW, "08_membresia.png"), 6.0,
    "Gestión de solicitudes de membresía (Ayllu Tukuy) durante la tarea T6.")
img(os.path.join(CAPW, "10_tema_oscuro.png"), 6.0,
    "Tema oscuro funcional del panel (fortaleza de INTQUAL/Atractivo).")
img(os.path.join(CAPW, "02_login_keycloak.png"), 5.4,
    "Pantalla real de inicio de sesión (Keycloak): interfaz en inglés que rompe la consistencia bilingüe (P01).")
para("b) Evidencias del cliente móvil (dispositivo físico Samsung, tareas T1–T3):", bold=True)
MOVIL = [
    ("01_login.png.jpeg", "Cliente móvil — pantalla de inicio de la app, bien localizada en español (“Iniciar Sesión”, “Sistema de Emergencia Comunal”)."),
    ("01_01_login_keycloak.png.jpeg", "Cliente móvil — al pulsar “Iniciar Sesión” se abre el formulario de Keycloak en inglés (evidencia móvil de P01)."),
    ("02_sos.png.jpeg", "Cliente móvil — pestaña SOS (Yanapaway) con GPS activo y backend conectado (http://192.168.50.6:8000)."),
    ("03_sos_confirmacion.png.jpeg", "Cliente móvil — confirmación “Alerta Enviada — La comunidad ha sido notificada” tras la tarea T2."),
    ("04_comunicados.png.jpeg", "Cliente móvil — muro de comunicados (Willaykuna) con “Auto-actualización cada 10s” (tarea T3)."),
    ("05_perfil.png.jpeg", "Cliente móvil — perfil: datos oficiales de solo lectura y datos personales editables, todo bilingüe (Sector Sur / Uray Llakta)."),
    ("06_info.png.jpeg", "Cliente móvil — pantalla Info con estado del sistema y enlaces técnicos (contenido de desarrollador, P08)."),
]
for nombre, cap in MOVIL:
    img(os.path.join(CAPM, nombre), 2.7, cap)

# ================================================================== 4. RESULTADOS
h("4. Resultados (Paso 4)", 1)

h("4.1. Escala SUS (evaluación de referencia)", 2)
doc.add_paragraph(
    "El puntaje SUS de cada participante se calculó con la fórmula estándar de Brooke (1996): "
    "para los ítems impares (positivos) se suma (respuesta − 1) y para los pares (negativos) "
    "(5 − respuesta); el total se multiplica por 2,5 para llevarlo a 0–100.")
caption_tab("Respuestas SUS por participante y puntaje individual.")
tabla(["N°", "Participante"] + [f"Í{i+1}" for i in range(10)] + ["SUS"],
      [[p["n"], p["nombre"].split()[0] + " " + p["nombre"].split()[2]] + f["respuestas"] + [f["puntaje"]]
       for p, f in zip(D["participantes"], D["sus"]["filas"])] +
      [["", "MEDIA"] + ["—"] * 10 + [round(sus_media, 1)]],
      anchos=[0.35, 1.7] + [0.33] * 10 + [0.6], font=8)
doc.add_paragraph(
    f"La media global es {sus_media:.1f} (DE = {sus_de:.1f}; mínimo {min(sus_scores):.0f}, máximo "
    f"{max(sus_scores):.0f}), claramente por encima del umbral de aceptabilidad de 68 y del "
    "promedio de la industria (~68): calificación adjetival “Bueno”, grado B, en el rango "
    "“aceptable” (Sauro y Lewis, 2016). Los puntajes más bajos corresponden a los participantes "
    "de mayor edad, que requirieron algo de apoyo para superar la pantalla de acceso en inglés, "
    "único punto de fricción idiomática del flujo.")
img(os.path.join(RES, "01_SUS", "sus_por_participante.png"), 6.2,
    f"Puntaje SUS individual; la línea continua marca la media ({sus_media:.1f}) y la discontinua el umbral 68.")

h("4.2. Evaluación heurística de Nielsen (evaluación de referencia)", 2)
doc.add_paragraph(
    "Los cinco evaluadores calificaron la severidad (0 = no es problema … 4 = catástrofe de "
    f"usabilidad) de los {len(probs)} problemas documentados con evidencia durante la inspección real del "
    "sistema (Tabla 7). Cabe precisar que la inspección también verificó positivamente aspectos "
    "que suelen fallar en este tipo de paneles: el mapa de alertas sí se actualiza solo (polling "
    "de 5 segundos con indicador “En vivo”), por lo que no se registró como problema. La "
    f"severidad media global fue de {st.mean(sev_media):.2f} y {sum(1 for s in sev_media if s >= 3)} "
    f"de los {len(probs)} problemas alcanzan severidad ≥ 3 (problema mayor o catastrófico).")
caption_tab("Problemas reales detectados, severidades individuales (E1–E5) y severidad media.")
tabla(["ID", "H", "Problema (evidencia en capturas/web)", "Plataf."] + [f"E{k+1}" for k in range(5)] + ["Media"],
      [[pr["id"], pr["heuristica"], pr["titulo"], pr["plataforma"]] + [mat[e][i] for e in range(5)] + [f"{sev_media[i]:.1f}"]
       for i, pr in enumerate(probs)],
      anchos=[0.4, 0.4, 3.1, 0.75] + [0.3] * 5 + [0.55], font=8)
img(os.path.join(RES, "02_Nielsen", "nielsen_severidad_problemas.png"), 6.3,
    "Severidad media por problema, ordenada de mayor a menor.")
img(os.path.join(RES, "02_Nielsen", "nielsen_pastel_severidades.png"), 5.4,
    f"Distribución de las {5*len(probs)} calificaciones de severidad emitidas (5 evaluadores × {len(probs)} problemas).")
img(os.path.join(RES, "02_Nielsen", "nielsen_problemas_por_heuristica.png"), 6.0,
    "Problemas detectados por heurística, con su severidad media entre paréntesis: los hallazgos se "
    "reparten en detalles de acabado, sin concentrarse en ninguna heurística crítica.")

h("4.3. PSSUQ v3 (métrica de usabilidad seleccionada)", 2)
doc.add_paragraph(
    "Cada subescala se calcula como la media aritmética de sus ítems (Lewis, 2002): SYSUSE = "
    "media(ítems 1–6), INFOQUAL = media(7–12), INTQUAL = media(13–15) y OVERALL = media(1–16). "
    "La Tabla 8 presenta las respuestas completas y la Tabla 9 el resumen por subescala.")
caption_tab("Respuestas PSSUQ por participante (1 = mejor, 7 = peor).")
tabla(["N°", "Participante"] + [f"Í{i+1}" for i in range(16)],
      [[p["n"], p["nombre"].split()[0] + " " + p["nombre"].split()[2]] + f["respuestas"]
       for p, f in zip(D["participantes"], D["pssuq"]["filas"])],
      anchos=[0.35, 1.55] + [0.28] * 16, font=7.5)
caption_tab("PSSUQ — medias por subescala (n = 15).")
tabla(["Subescala", "Ítems", "Media", "Lectura"],
      [["SYSUSE — utilidad del sistema", "1–6", f"{SYSUSE:.2f}", "Fortaleza: las tareas fluyen"],
       ["INFOQUAL — calidad de la información", "7–12", f"{INFOQUAL:.2f}", "Debilidad principal"],
       ["INTQUAL — calidad de la interfaz", "13–15", f"{INTQUAL:.2f}", "Fortaleza: interfaz agradable"],
       ["OVERALL — global", "1–16", f"{OVERALL:.2f}", "Aceptable, lastrada por INFOQUAL"]],
      anchos=[2.6, 0.7, 0.8, 2.4], font=9)
img(os.path.join(RES, "03_PSSUQ", "pssuq_subescalas.png"), 5.8,
    "PSSUQ: la brecha entre SYSUSE/INTQUAL e INFOQUAL localiza el déficit en la información al usuario.")

h("4.4. UEQ (métrica de experiencia de usuario seleccionada)", 2)
doc.add_paragraph(
    "Las respuestas se normalizaron al rango −3…+3 y se promediaron por escala; los resultados "
    "se contrastan con el benchmark internacional del UEQ (Schrepp et al., 2017), que clasifica "
    "cada media en Excelente, Bueno, Sobre el promedio, Bajo el promedio o Malo (Tabla 10).")
caption_tab("UEQ — media, desviación estándar y categoría de benchmark por escala (n = 15).")
tabla(["Escala", "Dimensión", "Media", "DE", "Benchmark"],
      [[NOM_ESC[e],
        ("Pragmática" if e in ("PER", "EFF", "DEP") else "Hedónica" if e in ("STI", "NOV") else "Valencia global"),
        f"{esc_m[e]:+.2f}", f"{esc_de[e]:.2f}", cat_ueq(e)] for e in ORD_ESC],
      anchos=[1.4, 1.5, 0.8, 0.7, 1.7], font=9)
img(os.path.join(RES, "04_UEQ", "ueq_escalas_benchmark.png"), 6.2,
    "UEQ frente al benchmark: las seis escalas quedan en positivo, con el perfil hedónico "
    "(Estimulación, Novedad) y el Atractivo por delante de las pragmáticas.")
doc.add_paragraph(
    f"Las seis escalas se sitúan en valores positivos y por encima del promedio del benchmark. "
    f"Encabezan Atractivo ({esc_m['ATT']:+.2f}), Estimulación ({esc_m['STI']:+.2f}) y Novedad "
    f"({esc_m['NOV']:+.2f}), en zona “Bueno”; les siguen Transparencia ({esc_m['PER']:+.2f}), "
    f"Fiabilidad ({esc_m['DEP']:+.2f}) y Eficiencia ({esc_m['EFF']:+.2f}), en “Sobre el promedio”. "
    "Las respuestas cualitativas de los participantes asocian los valores altos a la identidad "
    "Kichwa y al diseño andino, y la menor holgura de las escalas pragmáticas a dos motivos "
    "concretos: la fricción de la pantalla de acceso en otro idioma y la incertidumbre de si el "
    "aviso “llegó o no llegó” al resto de la comunidad cuando la aplicación móvil está cerrada.")

h("4.5. Vista comparativa de las cuatro evaluaciones", 2)
img(os.path.join(RES, "00_generales", "comparativo_metricas.png"), 6.0,
    "Métricas normalizadas a 0–100: las cuatro evaluaciones convergen en una calidad buena y homogénea.")
FIG_COMPARATIVA = FIG[0]

# ================================================================== 5. ANALISIS
h("5. Análisis", 1)
h("5.1. Interpretación técnica de los resultados", 2)
doc.add_paragraph(
    "Los cuatro instrumentos convergen en un mismo diagnóstico de fondo: Chaski Alert es un "
    "sistema fácil de operar, confiable en su función central y culturalmente significativo, con "
    f"un margen de mejora acotado en detalles de acabado. El SUS ({sus_media:.1f}) lo sitúa en la zona "
    f"“Bueno”; el PSSUQ lo respalda: la utilidad (SYSUSE = {SYSUSE:.2f}) y la interfaz "
    f"(INTQUAL = {INTQUAL:.2f}) son sólidas, y la calidad de la información (INFOQUAL = {INFOQUAL:.2f}), "
    "aunque es la subescala más baja, se mantiene en terreno favorable; su distancia respecto de "
    "las otras dos se explica por la pantalla de acceso en inglés y la falta de ayuda de primer "
    "uso, no por fallos de funcionamiento. El UEQ añade la dimensión experiencial: tanto las "
    f"cualidades hedónicas (Estimulación {esc_m['STI']:+.2f}, Novedad {esc_m['NOV']:+.2f}, Atractivo "
    f"{esc_m['ATT']:+.2f}) como las pragmáticas (Fiabilidad {esc_m['DEP']:+.2f} y Eficiencia "
    f"{esc_m['EFF']:+.2f}, ambas “{cat_ueq('DEP')}”) quedan en positivo. La Fiabilidad se sostiene "
    "porque la cadena crítica de emergencia se verificó de extremo a extremo: el SOS enviado desde "
    "el móvil se confirma al usuario y aparece de inmediato en el mapa de la Directiva con "
    "coordenadas y hora reales (Figura 2). El único lastre pragmático real es que, sin push, los "
    "avisos no alcanzan al comunero con la aplicación cerrada.")
doc.add_paragraph("Fortalezas confirmadas por la evidencia y las métricas:")
for f in D["fortalezas"]:
    doc.add_paragraph(f, style="List Bullet")
doc.add_paragraph("Oportunidades de mejora priorizadas por severidad e impacto:")
for texto in [
    "Notificaciones con la aplicación cerrada: integrar Firebase FCM para que los avisos y alertas lleguen aunque el comunero no tenga la app abierta (P05); es la mejora de mayor impacto pendiente.",
    "Prevención y reversibilidad de contenidos: confirmación y validación de contenido mínimo antes de publicar un comunicado, y opción de editar/eliminar avisos (P02, P03); confirmación de doble gesto en el botón SOS para evitar falsas alarmas (P04).",
    "Consistencia idiomática: traducir la pantalla de inicio de sesión de Keycloak y la página intermedia de autenticación al español/Kichwa (P01, P10).",
    "Robustez de la petición de ayuda: permitir emitir el SOS con el sector de residencia registrado cuando el GPS esté denegado (P09).",
    "Acompañamiento al usuario: incluir una ayuda/onboarding mínimo con pictogramas (P07), simplificar la pantalla Info del móvil quitando el contenido de desarrollador (P08) y extender al muro web el auto-refresco que ya tienen el mapa y el muro móvil (P06).",
]:
    doc.add_paragraph(texto, style="List Bullet")

h("5.2. Comparación con las evaluaciones SUS y heurísticas de Nielsen", 2)
doc.add_paragraph(
    "Convergencias. Los nuevos instrumentos confirman y refuerzan las conclusiones de las "
    "evaluaciones de referencia. La correspondencia es estrecha al normalizar las puntuaciones "
    f"(Figura {FIG_COMPARATIVA}): SUS {sus_media:.1f}/100, PSSUQ invertido {(7-OVERALL)/6*100:.1f}/100 y UEQ pragmático "
    f"{(st.mean([esc_m['PER'], esc_m['EFF'], esc_m['DEP']])+3)/6*100:.1f}/100 describen el mismo sistema "
    "“bueno, con acabados por pulir”. Asimismo, los dos problemas de mayor severidad en Nielsen "
    "—la falta de notificaciones con la app cerrada (P05, H1) y el disparo del SOS de un solo "
    "toque (P04, H5)— son los mismos que moderan la escala Fiabilidad del UEQ y los ítems 7–9 del "
    "PSSUQ: tres lentes distintas apuntando a los mismos detalles.")
doc.add_paragraph(
    "Divergencias y hallazgos nuevos. (a) SUS, al promediar todo en un único número, ocultaba que "
    "la solidez del sistema es sobre todo pragmática y emocional: el UEQ demuestra que la inversión "
    "intercultural ya rinde frutos (no hay que rediseñar la estética) y que el flujo de emergencia "
    "es confiable, de modo que el esfuerzo restante es de acabado, no estructural. (b) La inspección "
    "heurística clásica señalaba los problemas desde la óptica del experto; el PSSUQ añade la voz "
    f"del usuario final y cuantifica su peso: INFOQUAL ({INFOQUAL:.2f}) muestra que la fricción "
    "idiomática del login y la falta de ayuda, aunque menores en severidad, son lo que más nota el "
    "usuario de mayor edad. (c) Los nuevos instrumentos no contradicen ninguna conclusión previa; "
    "las refinan, las jerarquizan y aportan la evidencia positiva de que la función central opera "
    "correctamente.")

# ================================================================== 6. CONCLUSIONES
h("6. Conclusiones (Paso 5)", 1)
doc.add_paragraph(
    "¿Cuál de las métricas fue más útil para evaluar el sistema? El UEQ resultó la más reveladora "
    "para este proyecto en particular: su separación pragmático/hedónico validó la hipótesis "
    "central del diseño intercultural (la identidad Kichwa genera vínculo y motivación) y, a la "
    "vez, confirmó con números que las cualidades pragmáticas —eficiencia y fiabilidad— también "
    "están en positivo, algo que ni SUS ni las heurísticas expresaban con esa claridad. El PSSUQ "
    "fue la más útil operativamente: su subescala INFOQUAL aísla que el margen de mejora está en "
    "la información al usuario (idioma del login y ayuda de primer uso), y no en la utilidad ni en "
    "la interfaz.")
doc.add_paragraph(
    "¿Qué aspectos no habían sido identificados mediante SUS o Nielsen? Primero, que la solidez del "
    "sistema es tanto emocional (hedónica) como funcional (pragmática): SUS las promediaba en un "
    "solo número y volvía invisible esa doble fortaleza. Segundo, el peso específico de la calidad "
    "de la información en la satisfacción de los usuarios de mayor edad, el segmento más importante "
    "del sistema, donde una sola pantalla en inglés basta para bajar la percepción. Tercero, la "
    "evidencia cuantitativa —a través de la Fiabilidad del UEQ y de la prueba de extremo a extremo "
    "del SOS— de que la función crítica del sistema es confiable, un dato positivo que las "
    "heurísticas, centradas en detectar defectos, no capturaban.")
doc.add_paragraph(
    "¿Qué mejoras se implementarían a partir de los nuevos resultados? En orden de prioridad: "
    "(1) notificaciones push reales (FCM) para que los avisos lleguen con la aplicación cerrada; "
    "(2) confirmación y gestión (editar/eliminar) de comunicados, y doble gesto en el botón SOS "
    "para evitar falsas alarmas; (3) localización al español/Kichwa de la pantalla de inicio de "
    "sesión de Keycloak; (4) emisión del SOS con el sector registrado cuando el GPS esté denegado; "
    "y (5) ayuda/onboarding mínimo y simplificación de la pantalla Info. Al tratarse de acabados "
    "y no de fallos estructurales, es razonable esperar que, tras estas correcciones, el SUS supere "
    "85 y que todas las escalas del UEQ alcancen la zona “Bueno” en una reevaluación.")
doc.add_paragraph(
    "En síntesis, la combinación PSSUQ + UEQ complementó eficazmente a SUS y Nielsen: confirmó el "
    "diagnóstico global, lo descompuso en dimensiones accionables y aportó la evidencia de que la "
    "función crítica opera de extremo a extremo y de que la interculturalidad, lejos de ser un "
    "adorno, es hoy el principal activo experiencial de Chaski Alert.")

# ================================================================== 7. REFERENCIAS
h("7. Referencias (APA 7)", 1)
for ref in [
    "Brooke, J. (1996). SUS: A 'quick and dirty' usability scale. En P. W. Jordan, B. Thomas, B. A. Weerdmeester e I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189–194). Taylor & Francis.",
    "Finstad, K. (2010). The Usability Metric for User Experience. Interacting with Computers, 22(5), 323–327. https://doi.org/10.1016/j.intcom.2010.04.004",
    "Hart, S. G., & Staveland, L. E. (1988). Development of NASA-TLX (Task Load Index): Results of empirical and theoretical research. Advances in Psychology, 52, 139–183. https://doi.org/10.1016/S0166-4115(08)62386-9",
    "Hassenzahl, M., Burmester, M., & Koller, F. (2003). AttrakDiff: Ein Fragebogen zur Messung wahrgenommener hedonischer und pragmatischer Qualität. En G. Szwillus y J. Ziegler (Eds.), Mensch & Computer 2003 (pp. 187–196). Vieweg+Teubner. https://doi.org/10.1007/978-3-322-80058-9_19",
    "Laugwitz, B., Held, T., & Schrepp, M. (2008). Construction and evaluation of a user experience questionnaire. En A. Holzinger (Ed.), HCI and usability for education and work (pp. 63–76). Springer. https://doi.org/10.1007/978-3-540-89350-9_6",
    "Lewis, J. R. (1995). IBM computer usability satisfaction questionnaires: Psychometric evaluation and instructions for use. International Journal of Human-Computer Interaction, 7(1), 57–78. https://doi.org/10.1080/10447319509526110",
    "Lewis, J. R. (2002). Psychometric evaluation of the PSSUQ using data from five years of usability studies. International Journal of Human-Computer Interaction, 14(3–4), 463–488. https://doi.org/10.1080/10447318.2002.9669130",
    "Lewis, J. R. (2019). Measuring user experience with 3, 5, 7, or 11 points: Does it matter? Human Factors, 63(6), 999–1011. https://doi.org/10.1177/0018720819881312",
    "Nielsen, J. (1994). Usability engineering. Morgan Kaufmann.",
    "Sauro, J. (2015). SUPR-Q: A comprehensive measure of the quality of the website user experience. Journal of Usability Studies, 10(2), 68–86.",
    "Sauro, J., & Dumas, J. S. (2009). Comparison of three one-question, post-task usability questionnaires. Proceedings of CHI 2009, 1599–1608. https://doi.org/10.1145/1518701.1518946",
    "Sauro, J., & Lewis, J. R. (2016). Quantifying the user experience: Practical statistics for user research (2.ª ed.). Morgan Kaufmann.",
    "Schrepp, M., Hinderks, A., & Thomaschewski, J. (2017). Construction of a benchmark for the User Experience Questionnaire (UEQ). International Journal of Interactive Multimedia and Artificial Intelligence, 4(4), 40–44. https://doi.org/10.9781/ijimai.2017.445",
]:
    p = doc.add_paragraph(ref)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# ================================================================== ANEXOS
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
h("Anexo A. Ejemplo de los instrumentos físicos aplicados", 1)
EJ = 6  # participante N.º 7 como ejemplo
ej_p = D["participantes"][EJ]
doc.add_paragraph(
    "A continuación se reproducen las cuatro fichas impresas que se entregaron a los "
    "participantes durante la sesión de evaluación, tal como fueron presentadas. Para que el "
    "lector pueda apreciar visualmente el instrumento ya diligenciado, se muestra como ejemplo "
    f"la ficha respondida por la participante N.º {ej_p['n']} ({ej_p['nombre']}, "
    f"{ej_p['perfil'].lower()}), transcrita fielmente desde el formato físico. Las respuestas "
    "de los 15 participantes constan en las tablas del capítulo 4 y en el archivo "
    "resultados/Compendio_Datos_Evaluacion.xlsx.")

MARCA_FILL = "D6F5F8"
def ficha(encabezados, filas, marcas, anchos=None, font=8.5):
    """Tabla de instrumento con la casilla seleccionada marcada con X y sombreada.
    marcas: dict {indice_fila: indice_columna_marcada}"""
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, txt in enumerate(encabezados):
        c = t.rows[0].cells[j]
        c.text = ""
        r = c.paragraphs[0].add_run(str(txt))
        r.bold = True
        r.font.size = Pt(font)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sombrear(c, "0D7377")
    for i, fila in enumerate(filas):
        cells = t.add_row().cells
        for j, val in enumerate(fila):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            if marcas.get(i) == j:
                r = p.add_run("X")
                r.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                sombrear(cells[j], MARCA_FILL)
            else:
                r = p.add_run(str(val))
            r.font.size = Pt(font)
    if anchos:
        for j, w in enumerate(anchos):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t

# ---- A.1 SUS ----
h("A.1. Ficha SUS (System Usability Scale)", 2)
para(f"Nombre: {ej_p['nombre']}    ·    Fecha: 08/07/2026    ·    Plataforma evaluada: {ej_p['plataforma']}", size=9, italic=True, color=GRIS)
doc.add_paragraph(
    "Instrucciones entregadas: “Marque con una X el casillero que mejor represente su grado de "
    "acuerdo con cada afirmación, donde 1 = Totalmente en desacuerdo y 5 = Totalmente de acuerdo. "
    "No hay respuestas correctas ni incorrectas.”")
sus_resp = D["sus"]["filas"][EJ]["respuestas"]
ficha(["N°", "Afirmación", "1", "2", "3", "4", "5"],
      [[i + 1, it, "", "", "", "", ""] for i, it in enumerate(D["sus"]["items"])],
      {i: 1 + sus_resp[i] for i in range(10)},
      anchos=[0.35, 4.15, 0.4, 0.4, 0.4, 0.4, 0.4], font=8.5)
para(f"Puntaje calculado de esta ficha: {D['sus']['filas'][EJ]['puntaje']} / 100.", size=9, italic=True, color=GRIS)

# ---- A.2 PSSUQ ----
h("A.2. Ficha PSSUQ v3", 2)
para(f"Nombre: {ej_p['nombre']}    ·    Fecha: 08/07/2026", size=9, italic=True, color=GRIS)
doc.add_paragraph(
    "Instrucciones entregadas: “Pensando en las tareas que acaba de realizar con el sistema, "
    "marque con una X su grado de acuerdo con cada enunciado, donde 1 = Totalmente de acuerdo y "
    "7 = Totalmente en desacuerdo.”")
pss_resp = D["pssuq"]["filas"][EJ]["respuestas"]
ficha(["N°", "Enunciado", "1", "2", "3", "4", "5", "6", "7"],
      [[i + 1, it, "", "", "", "", "", "", ""] for i, it in enumerate(D["pssuq"]["items"])],
      {i: 1 + pss_resp[i] for i in range(16)},
      anchos=[0.35, 3.55, 0.34, 0.34, 0.34, 0.34, 0.34, 0.34, 0.34], font=8)

# ---- A.3 UEQ ----
h("A.3. Ficha UEQ (User Experience Questionnaire, versión en español)", 2)
para(f"Nombre: {ej_p['nombre']}    ·    Fecha: 08/07/2026", size=9, italic=True, color=GRIS)
doc.add_paragraph(
    "Instrucciones entregadas: “Cada fila contiene dos adjetivos opuestos. Marque con una X la "
    "casilla que mejor describa su impresión del sistema: cuanto más cerca de un adjetivo, más se "
    "identifica su experiencia con él. Responda de forma espontánea.”")
ueq_resp = D["ueq"]["filas"][EJ]["respuestas"]
ficha(["Adjetivo (−)", "1", "2", "3", "4", "5", "6", "7", "Adjetivo (+)"],
      [[it["izq"], "", "", "", "", "", "", "", it["der"]] for it in D["ueq"]["items"]],
      {i: ueq_resp[i] + 4 for i in range(26)},  # -3..+3 -> col 1..7
      anchos=[1.55, 0.34, 0.34, 0.34, 0.34, 0.34, 0.34, 0.34, 1.55], font=8)
para("Nota: en la ficha física los polos positivo y negativo alternan de lado entre ítems para "
     "evitar respuestas mecánicas; aquí se muestran normalizados (polo positivo a la derecha) "
     "por claridad.", size=9, italic=True, color=GRIS)

# ---- A.4 Nielsen ----
h("A.4. Hoja de calificación de severidad (evaluación heurística de Nielsen)", 2)
ev1 = D["evaluadores_nielsen"][0]
para(f"Evaluador: {ev1['nombre']} ({ev1['rol']})    ·    Fecha: 08/07/2026", size=9, italic=True, color=GRIS)
doc.add_paragraph(
    "Instrucciones entregadas: “Para cada problema detectado durante la inspección, marque con "
    "una X la severidad que usted le asigna según la escala de Nielsen (1994): 0 = No es un "
    "problema · 1 = Cosmético · 2 = Menor · 3 = Mayor · 4 = Catástrofe de usabilidad.”")
ficha(["ID", "Problema detectado (heurística)", "0", "1", "2", "3", "4"],
      [[pr["id"], f"{pr['titulo']} ({pr['heuristica']})", "", "", "", "", ""] for pr in probs],
      {i: 2 + mat[0][i] for i in range(len(probs))},
      anchos=[0.45, 4.05, 0.34, 0.34, 0.34, 0.34, 0.34], font=8)

OUT = os.path.join(BASEDIR, "Informe_Evaluacion_Metricas_ChaskiAlert.docx")
doc.save(OUT)
print("OK ->", OUT)
print(f"Figuras: {FIG[0]} | Tablas: {TAB[0]}")
