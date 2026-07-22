import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

base_dir = r"c:\Users\trejo\Desktop\S9\Inter\P_I\PR\INTER_PROYECTO\DOCUMENTACION_ISO_12207\5._Pruebas"
target_dir = r"c:\Users\trejo\Desktop\S9\Inter\P_I\PR\INTER_PROYECTO\Tarea_22_07\5. Pruebas-20260722T071144Z-1-001\5. Pruebas"
os.makedirs(base_dir, exist_ok=True)
os.makedirs(target_dir, exist_ok=True)

md_content = """# VV-03 Informe y Resultados de Pruebas

**Estándar ISO/IEC 12207 — Proceso de Evaluación de Calidad**  
**Proyecto:** CHASKI ALERT — Sistema de Alerta Comunitaria Intercultural  
**Código de Documento:** Doc-RP-001  
**Versión:** 1.5  
**Fecha de Emisión:** 22/07/2026  
**Elaborado por:** Alejandro Andrade, Milena Maldonado, Allan Panchi, Alex Trejo  

---

## 1. Resumen Ejecutivo
Se ejecutó un entorno de evaluación integral compuesto por **104 pruebas unitarias automatizadas** y **25 casos de prueba funcionales/heurísticos**, abarcando los módulos de autenticación, alertas SOS georreferenciadas, contingencia SMS, comunicados comunitarios, gestión de membresías y correcciones de usabilidad.

- **Total de Pruebas Unitarias Ejecutadas**: 104 (53 Backend, 22 Web, 29 Móvil)
- **Total de Casos Funcionales Ejecutados**: 25
- **Casos Exitosos (Aprobados)**: 100%
- **Casos Fallidos Pendientes**: 0%
- **Porcentaje de Éxito Global**: **100%**

---

## 2. Resultados de Pruebas Unitarias y Defectos Encontrados

| Aplicación | Herramienta | Pruebas | Tiempo | Estado |
|---|---|---|---|---|
| Backend (FastAPI) | pytest 9.1 | 53 | 1.3 s | 100% Aprobado |
| Panel web (Next.js) | Jest 30 + Testing Library | 22 | 1.3 s | 100% Aprobado |
| App móvil (Expo) | jest-expo 54 | 29 | 21 s | 100% Aprobado |
| **Total** | | **104** | | **100% Aprobado** |

*Nota: Todas pasan de forma aislada sin necesidad de levantar Docker, base de datos ni Keycloak, confirmando su naturaleza puramente unitaria.*

### Defecto crítico resuelto durante las pruebas
Durante las pruebas de la función `timeSince()` en el frontend web, se encontró un defecto manejando husos horarios negativos. La comprobación original `dateStr.includes("+") || dateStr.includes("Z")` fallaba al recibir `-05:00`, añadiendo un doble offset que mostraba `hace NaNd` en las tarjetas. La prueba forzó la corrección mediante validación Regex `/(?:Z|[+-]\d{2}:?\d{2})$/`, protegiendo al sistema ante una posible migración de la DB a `TIMESTAMPTZ`.

---

## 3. Resultados de Evaluaciones Estandarizadas de UX y Usabilidad

### 3.1 Escala SUS (System Usability Scale)
- **Muestra**: 15 Participantes externos (comuneros, agricultores, dirigentes, estudiantes).
- **Resultado Global SUS**: **80.2 / 100** (DE = 10.5).
- **Calificación Adjetival**: "Bueno" (Grado B, Rango Aceptable).

### 3.2 PSSUQ v3 (Post-Study System Usability Questionnaire)
*Escala de 1 a 7 (menor es mejor).*
- `SYSUSE` (Utilidad del Sistema): **2.18** (Fortaleza: flujo de tareas fluido).
- `INFOQUAL` (Calidad de la Información): **3.23** (Mejorada tras traducir Keycloak a Español).
- `INTQUAL` (Calidad de la Interfaz): **2.49** (Fortaleza: interfaz andina agradable).
- `OVERALL` (Puntuación Global): **2.62** (Aceptable y satisfactoria).

### 3.3 UEQ (User Experience Questionnaire)
*Escala normalizada de -3 a +3.*
- **Atractivo**: **+1.66** (Excelente / Hedónico Fuerte)
- **Novedad**: **+1.23** (Bueno - Pertinencia Kichwa valorada positivamente)
- **Eficiencia**: **+1.13** (Sobre el promedio)

---

## 4. Resumen de Ejecución de Casos de Prueba (QA Funcional)

| Módulo | Casos Evaluados | Exitosos | Fallidos | Estado Final |
|---|---|---|---|---|
| Autenticación Keycloak (OAuth2) | 3 | 3 | 0 | APROBADO |
| Emergencias SOS & Contingencia SMS | 4 | 4 | 0 | APROBADO |
| Mapa de Alertas Web (Leaflet) | 2 | 2 | 0 | APROBADO |
| Comunicados (Willaykuna & SWR) | 6 | 6 | 0 | APROBADO |
| Gestión Comunal (Ayllu) | 2 | 2 | 0 | APROBADO |
| Correcciones Heurísticas (P01-P10) | 8 | 8 | 0 | APROBADO |
| **TOTAL** | **25** | **25** | **0** | **100% EXITOSO** |

---

## 5. Responsables

A continuación, se detallan los responsables de este documento, sus roles dentro del proyecto y las funciones específicas asignadas en el marco del Sistema de Gestión de Calidad:

**Nombre:** Milena Maldonado  
**Rol:** Parte del equipo de desarrollo / Analista  
**Categoría profesional:** Ing. Software  
**Responsabilidad:** Verificar que la planificación cumpla con los estándares de calidad establecidos, supervisar el cumplimiento de los procesos definidos y proponer mejoras continuas en la documentación del proyecto.  
**Información de Contacto:** mvmaldonado3@espe.edu.ec  

**Nombre:** Alex Trejo  
**Rol / Cargo:** Líder de Proyecto / Parte del equipo de desarrollo  
**Categoría profesional:** Ing. Software  
**Responsabilidad:** Coordinar de manera general la planificación del proyecto, organizar las actividades por etapas, controlar su ejecución y asegurar el cumplimiento del cronograma y objetivos establecidos.  
**Información de Contacto:** aftrejo@espe.edu.ec  

**Nombre:** Alejandro Andrade  
**Rol:** Backend Developer / Database Developer  
**Categoría profesional:** Ing. Software  
**Responsabilidad:** Desarrollo de soluciones backend, arquitectura de sistema y control de bases de datos.  
**Información de Contacto:** laandrade9@espe.edu.ec  

**Nombre:** Allan Panchi  
**Rol:** Frontend Developer  
**Categoría profesional:** Ing. Software  
**Responsabilidad:** Desarrollo integral de Frontend, revisión, planificación y diseño de interfaces para usuario; siguiendo lineamientos y demás reglas usabilidad.  
**Información de Contacto:** avpanchi@espe.edu.ec  

---

## 6. Control de Cambios

| Versión | Fecha | Descripción del Cambio | Responsable |
|---|---|---|---|
| 1.0 | 13/08/2025 | Versión inicial de Resultados de Pruebas | Alejandro Andrade |
| 1.5 | 22/07/2026 | Inclusión de 104 métricas unitarias y resultados UX reales. | Alex Trejo |

---

________________________                             ___________________________________
 Director de Pruebas                                              Firma del Líder de Proyecto
 Alejandro Andrade                                                      Alex Trejo

---

## 7. Anexos
- **Repositorio de Github:** `https://github.com/Alex-Trejo/INTER_PROYECTO.git`
"""

with open(os.path.join(base_dir, "VV-03_Resultados_de_Pruebas.md"), "w", encoding="utf-8") as f:
    f.write(md_content)
with open(os.path.join(target_dir, "Resultados de Pruebas.md"), "w", encoding="utf-8") as f:
    f.write(md_content)

def set_cell_bg(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def create_resultados_docx(out_path):
    doc = Document()
    
    title = doc.add_heading('VV-03 Informe y Resultados de Pruebas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Estándar ISO/IEC 12207 — Proceso de Evaluación de Calidad').bold = True
    doc.add_paragraph('Proyecto: CHASKI ALERT — Sistema de Alerta Comunitaria Intercultural')
    doc.add_paragraph('Código: Doc-RP-001 | Versión: 1.5 | Fecha de Emisión: 22/07/2026')
    doc.add_paragraph('Elaborado por: Alejandro Andrade, Milena Maldonado, Allan Panchi, Alex Trejo')
    
    doc.add_heading('1. Resumen Ejecutivo', level=1)
    doc.add_paragraph('Se ejecutó un entorno de evaluación integral compuesto por 104 pruebas unitarias automatizadas y 25 casos de prueba funcionales/heurísticos, abarcando los módulos de autenticación, alertas SOS georreferenciadas, contingencia SMS, comunicados comunitarios, gestión de membresías y correcciones de usabilidad.')
    doc.add_paragraph('- Total de Pruebas Unitarias Ejecutadas: 104 (53 Backend, 22 Web, 29 Móvil)\n- Total de Casos Funcionales Ejecutados: 25\n- Casos Exitosos (Aprobados): 100%\n- Casos Fallidos Pendientes: 0%\n- Porcentaje de Éxito Global: 100%')

    doc.add_heading('2. Resultados de Pruebas Unitarias y Defectos Encontrados', level=1)
    t_unit = doc.add_table(rows=5, cols=5)
    t_unit.style = 'Table Grid'
    t_unit.rows[0].cells[0].text = 'Aplicación'
    t_unit.rows[0].cells[1].text = 'Herramienta'
    t_unit.rows[0].cells[2].text = 'Pruebas'
    t_unit.rows[0].cells[3].text = 'Tiempo'
    t_unit.rows[0].cells[4].text = 'Estado'
    unit_data = [
        ('Backend (FastAPI)', 'pytest 9.1', '53', '1.3 s', '100% Aprobado'),
        ('Panel web (Next.js)', 'Jest 30 + Testing Library', '22', '1.3 s', '100% Aprobado'),
        ('App móvil (Expo)', 'jest-expo 54', '29', '21 s', '100% Aprobado'),
        ('Total', '', '104', '', '100% Aprobado')
    ]
    for i, (a, h, p, t, e) in enumerate(unit_data):
        t_unit.rows[i+1].cells[0].text = a
        t_unit.rows[i+1].cells[1].text = h
        t_unit.rows[i+1].cells[2].text = p
        t_unit.rows[i+1].cells[3].text = t
        t_unit.rows[i+1].cells[4].text = e
    doc.add_paragraph('\nDefecto crítico resuelto durante las pruebas:\nDurante las pruebas de la función timeSince() en el frontend web, se encontró un defecto manejando husos horarios negativos (-05:00). La prueba detectó un doble offset (NaNd). Se corrigió con la Regex apropiada (/(?:Z|[+-]\d{2}:?\d{2})$/), protegiendo futuras migraciones de DB.')

    doc.add_heading('3. Resultados de Evaluaciones Estandarizadas de UX y Usabilidad', level=1)
    doc.add_heading('3.1 Escala SUS', level=2)
    doc.add_paragraph('- Muestra: 15 Participantes externos.\n- Resultado Global SUS: 80.2 / 100.\n- Calificación Adjetival: "Bueno" (Grado B, Rango Aceptable).')
    doc.add_heading('3.2 PSSUQ v3', level=2)
    doc.add_paragraph('- SYSUSE: 2.18 (Fortaleza: flujo de tareas fluido).\n- INFOQUAL: 3.23 (Mejorada tras traducir Keycloak a Español).\n- INTQUAL: 2.49 (Fortaleza: interfaz andina agradable).\n- OVERALL: 2.62 (Aceptable y satisfactoria).')
    doc.add_heading('3.3 UEQ', level=2)
    doc.add_paragraph('- Atractivo: +1.66 (Excelente)\n- Novedad: +1.23 (Bueno - Pertinencia Kichwa)\n- Eficiencia: +1.13 (Sobre el promedio)')

    doc.add_heading('4. Resumen de Ejecución de Casos de Prueba (QA Funcional)', level=1)
    t_qa = doc.add_table(rows=8, cols=5)
    t_qa.style = 'Table Grid'
    t_qa.rows[0].cells[0].text = 'Módulo'
    t_qa.rows[0].cells[1].text = 'Casos Evaluados'
    t_qa.rows[0].cells[2].text = 'Exitosos'
    t_qa.rows[0].cells[3].text = 'Fallidos'
    t_qa.rows[0].cells[4].text = 'Estado Final'
    qa_data = [
        ('Autenticación Keycloak (OAuth2)', '3', '3', '0', 'APROBADO'),
        ('Emergencias SOS & Contingencia SMS', '4', '4', '0', 'APROBADO'),
        ('Mapa de Alertas Web (Leaflet)', '2', '2', '0', 'APROBADO'),
        ('Comunicados (Willaykuna & SWR)', '6', '6', '0', 'APROBADO'),
        ('Gestión Comunal (Ayllu)', '2', '2', '0', 'APROBADO'),
        ('Correcciones Heurísticas (P01-P10)', '8', '8', '0', 'APROBADO'),
        ('TOTAL', '25', '25', '0', '100% EXITOSO')
    ]
    for i, (m, ce, ex, fa, es) in enumerate(qa_data):
        t_qa.rows[i+1].cells[0].text = m
        t_qa.rows[i+1].cells[1].text = ce
        t_qa.rows[i+1].cells[2].text = ex
        t_qa.rows[i+1].cells[3].text = fa
        t_qa.rows[i+1].cells[4].text = es

    doc.add_heading('5. Responsables', level=1)
    responsables_data = [
        ("Milena Maldonado", "Parte del equipo de desarrollo / Analista", "Verificar que la planificación cumpla con los estándares de calidad establecidos, supervisar el cumplimiento de los procesos definidos y proponer mejoras continuas en la documentación del proyecto.", "mvmaldonado3@espe.edu.ec"),
        ("Alex Trejo", "Líder de Proyecto / Parte del equipo de desarrollo", "Coordinar de manera general la planificación del proyecto, organizar las actividades por etapas, controlar su ejecución y asegurar el cumplimiento del cronograma y objetivos establecidos.", "aftrejo@espe.edu.ec"),
        ("Alejandro Andrade", "Backend Developer / Database Developer", "Desarrollo de soluciones backend, arquitectura de sistema y control de bases de datos.", "laandrade9@espe.edu.ec"),
        ("Allan Panchi", "Frontend Developer", "Desarrollo integral de Frontend, revisión, planificación y diseño de interfaces para usuario; siguiendo lineamientos y demás reglas usabilidad.", "avpanchi@espe.edu.ec")
    ]
    
    blue_bg = '9CC2E5'
    for resp in responsables_data:
        t_resp = doc.add_table(rows=5, cols=2)
        t_resp.style = 'Table Grid'
        
        t_resp.rows[0].cells[0].text = 'Nombre'
        set_cell_bg(t_resp.rows[0].cells[0], blue_bg)
        t_resp.rows[0].cells[1].text = resp[0]
        
        t_resp.rows[1].cells[0].text = 'Rol / Cargo' if "Líder" in resp[1] else 'Rol'
        set_cell_bg(t_resp.rows[1].cells[0], blue_bg)
        t_resp.rows[1].cells[1].text = resp[1]
        
        t_resp.rows[2].cells[0].text = 'Categoría profesional'
        set_cell_bg(t_resp.rows[2].cells[0], blue_bg)
        t_resp.rows[2].cells[1].text = 'Ing. Software'
        
        t_resp.rows[3].cells[0].text = 'Responsabilidad'
        set_cell_bg(t_resp.rows[3].cells[0], blue_bg)
        t_resp.rows[3].cells[1].text = resp[2]
        
        t_resp.rows[4].cells[0].text = 'Información de Contacto'
        set_cell_bg(t_resp.rows[4].cells[0], blue_bg)
        t_resp.rows[4].cells[1].text = resp[3]
        
        doc.add_paragraph()

    doc.add_heading('6. Control de Cambios', level=1)
    t_ctrl = doc.add_table(rows=3, cols=4)
    t_ctrl.style = 'Table Grid'
    t_ctrl.rows[0].cells[0].text = 'Versión'
    t_ctrl.rows[0].cells[1].text = 'Fecha'
    t_ctrl.rows[0].cells[2].text = 'Descripción del Cambio'
    t_ctrl.rows[0].cells[3].text = 'Responsable'
    
    cambios = [
        ('1.0', '13/08/2025', 'Versión inicial de Resultados de Pruebas', 'Alejandro Andrade'),
        ('1.5', '22/07/2026', 'Inclusión de 104 métricas unitarias y resultados UX reales.', 'Alex Trejo')
    ]
    for i, c in enumerate(cambios):
        t_ctrl.rows[i+1].cells[0].text = c[0]
        t_ctrl.rows[i+1].cells[1].text = c[1]
        t_ctrl.rows[i+1].cells[2].text = c[2]
        t_ctrl.rows[i+1].cells[3].text = c[3]

    doc.add_paragraph()
    firmas = doc.add_paragraph()
    firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firmas.add_run('________________________                             ___________________________________\n')
    firmas.add_run('  Director de Pruebas                                        Firma del Líder de Proyecto\n')
    firmas.add_run('   Alejandro Andrade                                                  Alex Trejo')

    doc.add_heading('7. Anexos', level=1)
    doc.add_paragraph('Repositorio de Github\nEnlace:\nhttps://github.com/Alex-Trejo/INTER_PROYECTO.git')

    doc.save(out_path)

create_resultados_docx(os.path.join(base_dir, "VV-03_Resultados_de_Pruebas.docx"))
# Guardado con Sufijo para evitar error de Permission Denied por si está abierto
create_resultados_docx(os.path.join(target_dir, "Resultados de Pruebas Completo.docx"))
